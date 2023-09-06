from docx import Document

# # создание документа
# document = Document()
# открытие документа
document = Document('/home/esal/PycharmProjects/WEB_parsing/others/Список личного состава для схемы оповещения(3).docx')
text = []
for paragraph in document.paragraphs:
    text.append(paragraph.text)
print('\n\n'.join(text))